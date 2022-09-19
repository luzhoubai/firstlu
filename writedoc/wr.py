
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches


def w():

    my_doc = Document()
    my_doc.add_heading('word instruction',0)

    my_doc.add_paragraph('this is my new para')
    my_doc.add_paragraph('this is my another new para')
    my_para = my_doc.add_paragraph('this is my third new para')
    my_run = my_para.add_run('this is my new block')

    my_run.font.size = Pt(18)
    my_run.font.italic = True
    my_run.font.bold = True
    my_run.font.color.rgb = RGBColor(255,0,0)
    my_run.font.name = 'Times New Roman'

    my_doc.add_picture('销售额.png',width=Inches(5.5))
    table = my_doc.add_table(rows=1,cols=3)
    row_0 = table.rows[0].cells
    row_0[0].text = 'number'
    row_0[1].text = 'name'
    row_0[2].text = 'age'

    datas=(
        (1,'张三','22'),
        (2,'李四','21'),
        (3,'王五','23'))

    for id,name,age in datas:
        row_cells = table.add_row().cells   #表格增加一行
        row_cells[0].text = str(id)
        row_cells[1].text = name
        row_cells[2].text = age

    my_doc.save('my word.docx')